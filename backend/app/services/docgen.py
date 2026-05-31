"""
Генерация документов: .docx из python-docx (или docxtpl для шаблонов), .pdf через fpdf2.
Файлы загружаются в Яндекс Object Storage (S3) по ключу {order_id}/{filename}.

Стандарт оформления российских претензий/заявлений:
- Шапка (получатель + отправитель) — правый верхний угол (~85мм)
- Заголовок (ПРЕТЕНЗИЯ / ЖАЛОБА) — по центру
- Тело — от левого края
- Блок подписи — дата слева + линия/инициалы справа, расшифровка под ней справа
- Блок подписи не разрывается переносом страницы
"""

import asyncio
import io
import logging
import re
from datetime import datetime, UTC
from pathlib import Path

from docx import Document as DocxDocument

from app.core.constants import MONTHS_RU
from app.services.storage import upload_bytes

logger = logging.getLogger(__name__)

TEMPLATES_DIR = Path(__file__).parent.parent / "templates"

_SAFE_SITUATION_ID = re.compile(r"^[a-z_]{1,64}$")
_SAFE_ORDER_ID = re.compile(r"^[0-9a-f\-]{36}$")


def _sanitize_situation_id(situation_id: str) -> str:
    if not _SAFE_SITUATION_ID.match(situation_id):
        raise ValueError(f"Invalid situation_id: {situation_id!r}")
    return situation_id


def _sanitize_order_id(order_id: str) -> str:
    if not _SAFE_ORDER_ID.match(order_id):
        raise ValueError(f"Invalid order_id: {order_id!r}")
    return order_id


def _find_template(situation_id: str, form_data: dict) -> Path | None:
    """Look for a .docx template: first by subtype, then by situation_id."""
    subtype = (
        form_data.get("problem_type")
        or form_data.get("violation_type")
        or form_data.get("incident_type")
    )
    if subtype and _SAFE_SITUATION_ID.match(str(subtype)):
        path = TEMPLATES_DIR / situation_id / f"{subtype}.docx"
        if path.resolve().is_relative_to(TEMPLATES_DIR.resolve()) and path.exists():
            return path
    path = TEMPLATES_DIR / f"{situation_id}.docx"
    if path.exists():
        return path
    return None


def _render_template(template_path: Path, context: dict) -> bytes:
    """Render a docxtpl template with the given context."""
    from docxtpl import DocxTemplate

    tpl = DocxTemplate(str(template_path))
    tpl.render(context)
    buf = io.BytesIO()
    tpl.save(buf)
    return buf.getvalue()


def _text_to_docx(text: str) -> bytes:
    doc = DocxDocument()
    for paragraph in text.split("\n"):
        doc.add_paragraph(paragraph)
    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


def _build_header(header_fields: list, form_data: dict) -> list[str]:
    """Детерминированно строит строки шапки из header_fields конфига и form_data.

    Персональные данные берутся напрямую из form_data — в LLM не уходят.
    Пустые поля пропускаются.
    """
    lines: list[str] = []
    for hf in header_fields:
        field = hf.field if hasattr(hf, "field") else hf.get("field", "")
        label = hf.label if hasattr(hf, "label") else hf.get("label")
        prefix = hf.prefix if hasattr(hf, "prefix") else hf.get("prefix")

        value = str(form_data.get(field, "")).strip()
        if not value:
            continue

        if prefix:
            value = prefix + value
        if label:
            lines.append(label)
        lines.append(value)
    return lines


def _render_right_block(pdf: "FPDF", lines: list[str], line_h: float = 5.5) -> None:
    RIGHT_COL_X = 105.0
    RIGHT_COL_W = 85.0
    for line in lines:
        if not line.strip():
            pdf.ln(2)
        else:
            pdf.set_x(RIGHT_COL_X)
            pdf.multi_cell(RIGHT_COL_W, line_h, line.strip(), align="L")


def _split_last_line(lines: list[str]) -> tuple[list[str], str | None]:
    """Отделяет последнюю непустую строку, чтобы держать её вместе с блоком подписи."""
    for i in range(len(lines) - 1, -1, -1):
        if lines[i].strip():
            return lines[:i], lines[i]
    return lines, None


def _render_sig_block(pdf: "FPDF", line_h: float = 6.0, tail_line: str | None = None) -> None:
    """Блок подписи: дата-бланк слева + подпись/расшифровка-бланк справа. Всё от руки.

    tail_line — последний абзац тела. Рендерится атомарно вместе с блоком подписи,
    чтобы подпись не уезжала на новую страницу одна без текста (keep-with-last-line).
    """
    PAGE_W  = 210.0
    LEFT_M  = 25.0
    RIGHT_M = 20.0
    BODY_W  = PAGE_W - LEFT_M - RIGHT_M
    SIG_W   = 80.0
    BLOCK_H = 6 + 2 * line_h + 5  # ln(6) перед блоком + две строки + запас

    tail_h = 0.0
    if tail_line:
        tail_h = pdf.multi_cell(0, 6, tail_line, dry_run=True, output="HEIGHT") + 1

    if pdf.get_y() + tail_h + BLOCK_H > pdf.h - pdf.b_margin:
        pdf.add_page()

    if tail_line:
        pdf.multi_cell(0, 6, tail_line)
        pdf.ln(1)

    DATE_TEXT = "«___» _________________ 20___ г."
    SIG_TEXT  = "_________________ / _________________"

    pdf.ln(6)
    pdf.set_x(LEFT_M)
    pdf.cell(BODY_W - SIG_W, line_h, DATE_TEXT)
    pdf.set_x(LEFT_M + BODY_W - SIG_W)
    pdf.cell(SIG_W, line_h, SIG_TEXT)
    pdf.ln(line_h)


def _docx_to_pdf_legacy(docx_bytes: bytes) -> bytes:
    """Legacy: конвертирует готовый .docx (шаблонный путь) в PDF через fpdf.

    Используется только когда для ситуации есть .docx-шаблон.
    Структура документа берётся из docx как есть — без парсинга шапки.
    """
    import io as _io
    from docx import Document as DocxDoc
    from fpdf import FPDF

    doc = DocxDoc(_io.BytesIO(docx_bytes))
    lines = [re.sub(r'(?<!\-)\-\-(?!\-)', '—', p.text) for p in doc.paragraphs]

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.add_font("FreeSans", fname="/usr/share/fonts/truetype/freefont/FreeSans.ttf")
    pdf.add_font("FreeSans", style="B", fname="/usr/share/fonts/truetype/freefont/FreeSansBold.ttf")
    pdf.set_margins(left=25, top=25, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_font("FreeSans", size=11)

    head_lines, tail_line = _split_last_line(lines)
    for line in head_lines:
        if not line.strip():
            pdf.ln(3)
        else:
            pdf.multi_cell(0, 6, line)
            pdf.ln(1)

    _render_sig_block(pdf, tail_line=tail_line)
    return bytes(pdf.output())


def _render_pdf(header: list[str], title: str, body: str) -> bytes:
    """Рендерит PDF из готовых частей: шапка (детерминированная), заголовок, тело."""
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.add_font("FreeSans", fname="/usr/share/fonts/truetype/freefont/FreeSans.ttf")
    pdf.add_font(
        "FreeSans", style="B",
        fname="/usr/share/fonts/truetype/freefont/FreeSansBold.ttf",
    )
    pdf.set_margins(left=25, top=25, right=20)
    pdf.set_auto_page_break(auto=True, margin=20)
    pdf.set_font("FreeSans", size=11)

    if header:
        _render_right_block(pdf, header)
        pdf.ln(4)

    if title:
        pdf.set_font("FreeSans", style="B", size=12)
        pdf.multi_cell(0, 7, title, align="C")
        pdf.set_font("FreeSans", size=11)
        pdf.ln(4)

    lines = body.split("\n")
    head_lines, tail_line = _split_last_line(lines)
    for line in head_lines:
        if not line.strip():
            pdf.ln(3)
            continue
        pdf.multi_cell(0, 6, line)
        pdf.ln(1)

    _render_sig_block(pdf, tail_line=tail_line)

    return bytes(pdf.output())


def s3_key(order_id: str, filename: str) -> str:
    return f"{order_id}/{filename}"


def _instruction_to_pdf(content: str, legal_refs: list[dict]) -> bytes:
    from fpdf import FPDF

    pdf = FPDF(orientation="P", unit="mm", format="A4")
    pdf.add_page()
    pdf.add_font("FreeSans", fname="/usr/share/fonts/truetype/freefont/FreeSans.ttf")
    pdf.add_font("FreeSans", style="B", fname="/usr/share/fonts/truetype/freefont/FreeSansBold.ttf")
    pdf.set_margins(left=25, top=25, right=25)
    pdf.set_auto_page_break(auto=True, margin=20)

    pdf.set_font("FreeSans", style="B", size=13)
    pdf.multi_cell(0, 8, "Инструкция по подаче претензии")
    pdf.ln(4)

    # AI-generated content
    pdf.set_font("FreeSans", size=11)
    for para in content.split("\n"):
        if not para.strip():
            pdf.ln(3)
            continue
        pdf.multi_cell(0, 6, para)
        pdf.ln(1)

    if legal_refs:
        pdf.ln(6)
        pdf.set_draw_color(200, 200, 200)
        pdf.line(pdf.get_x(), pdf.get_y(), 185, pdf.get_y())
        pdf.ln(5)

        pdf.set_font("FreeSans", style="B", size=11)
        pdf.multi_cell(0, 6, "Ссылки на законы из вашего заявления:")
        pdf.ln(3)

        pdf.set_font("FreeSans", size=10)
        for ref in legal_refs:
            law = ref.get("law", "")
            url = ref.get("url", "")
            if not law:
                continue
            pdf.write(6, "• ")
            if url:
                pdf.set_text_color(0, 80, 200)
                pdf.write(6, law, link=url)
                pdf.set_text_color(0, 0, 0)
            else:
                pdf.write(6, law)
            pdf.ln(7)

    return bytes(pdf.output())


async def generate_instruction(
    order_id: str,
    situation_id: str,
    content: str,
    legal_refs: list[dict],
) -> str:
    """Build instruction PDF, upload to S3. Returns S3 key."""
    safe_sid = _sanitize_situation_id(situation_id)
    safe_oid = _sanitize_order_id(order_id)

    pdf_bytes = await asyncio.get_running_loop().run_in_executor(
        None, _instruction_to_pdf, content, legal_refs
    )

    date_str = datetime.now(UTC).strftime("%Y%m%d")
    filename = f"instrukciya_{safe_sid}_{date_str}.pdf"
    key = s3_key(safe_oid, filename)

    await upload_bytes(key, pdf_bytes)
    return key


async def generate_document(
    order_id: str,
    situation_id: str,
    body: str,
    header: list[str],
    title: str,
    form_data: dict | None = None,
) -> tuple[str, str]:
    """
    Builds .docx and .pdf, uploads to S3.

    body — только текст тела (от LLM), без шапки и заголовка.
    header — строки шапки, построенные детерминированно из form_data.
    title — заголовок документа (ПРЕТЕНЗИЯ / ЖАЛОБА / ЗАЯВЛЕНИЕ).

    Returns (docx_key, pdf_key).
    """
    safe_sid = _sanitize_situation_id(situation_id)
    safe_oid = _sanitize_order_id(order_id)

    fd = form_data or {}
    loop = asyncio.get_running_loop()

    template_path = _find_template(safe_sid, fd)
    if template_path:
        now = datetime.now(UTC)
        context = {
            **fd,
            "ai_narrative": body,
            "today": f"{now.day} {MONTHS_RU[now.month - 1]} {now.year} года",
        }
        docx_bytes = await loop.run_in_executor(None, _render_template, template_path, context)
        logger.info("Rendered template %s for order %s", template_path.name, safe_oid)
        pdf_bytes = await loop.run_in_executor(None, _docx_to_pdf_legacy, docx_bytes)
    else:
        full_text = "\n".join(header) + f"\n\n{title}\n\n" + body if header else f"{title}\n\n{body}"
        docx_bytes = await loop.run_in_executor(None, _text_to_docx, full_text)
        pdf_bytes = await loop.run_in_executor(None, _render_pdf, header, title, body)

    date_str = datetime.now(UTC).strftime("%Y%m%d")
    docx_key = s3_key(safe_oid, f"pretenziya_{safe_sid}_{date_str}.docx")
    pdf_key = s3_key(safe_oid, f"pretenziya_{safe_sid}_{date_str}.pdf")

    await asyncio.gather(
        upload_bytes(docx_key, docx_bytes),
        upload_bytes(pdf_key, pdf_bytes),
    )

    return docx_key, pdf_key
